import unittest
import tempfile
from pathlib import Path
from docx import Document as WordDocument

from backend.core.rag_pipeline import RAGPipeline
from backend.services.indexing_service import _attach_parent_blocks
from backend.services.query_rewriter import QueryRewriter
from backend.services.evaluation_service import retrieval_metrics
from backend.db.models import EvaluationItem
from backend.services.document_parser import DocumentParser


class RagFeatureTests(unittest.TestCase):
    def test_rrf_combines_rankings(self):
        vector = [{"chunk_id": "a", "search_type": "vector"}, {"chunk_id": "b", "search_type": "vector"}]
        keyword = [{"chunk_id": "b", "search_type": "keyword"}, {"chunk_id": "c", "search_type": "keyword"}]
        result = RAGPipeline._rrf_fuse([vector, keyword], 60)
        self.assertEqual(result[0]["chunk_id"], "b")
        self.assertEqual(result[0]["source_ranks"], {"vector": 2, "keyword": 1})

    def test_diversification_limits_one_document_monopoly(self):
        results = [
            {"chunk_id": "a1", "document_id": "a"},
            {"chunk_id": "a2", "document_id": "a"},
            {"chunk_id": "a3", "document_id": "a"},
            {"chunk_id": "b1", "document_id": "b"},
            {"chunk_id": "c1", "document_id": "c"},
        ]
        selected = RAGPipeline._diversify_results(results, 4, max_per_document=2)
        self.assertEqual([item["chunk_id"] for item in selected], ["a1", "a2", "b1", "c1"])

    def test_labelled_retrieval_metrics_use_correct_documents(self):
        item = EvaluationItem(
            question="test", expected_document_titles=["correct.docx", "second.docx"],
            expected_chunk_ids=[], expected_pages=[], expected_keywords=["ambiguous"],
        )
        results = [
            {"chunk_id": "x", "document_title": "wrong.docx", "content": "ambiguous", "metadata": {}},
            {"chunk_id": "y", "document_title": "correct.docx", "content": "answer", "metadata": {}},
            {"chunk_id": "z", "document_title": "second.docx", "content": "answer", "metadata": {}},
        ]
        metrics = retrieval_metrics(item, results)
        self.assertEqual(metrics["relevant_ranks"], [2, 3])
        self.assertEqual(metrics["document_recall_at_10"], 1.0)
        self.assertEqual(metrics["recall_at_5"], 1.0)
    def test_parent_blocks_respect_limit_for_normal_children(self):
        chunks = [{"content": "a" * 600, "metadata": {"page": 1}}, {"content": "b" * 600, "metadata": {"page": 1}}, {"content": "c" * 600, "metadata": {"page": 1}}]
        result = _attach_parent_blocks(chunks, 1200)
        self.assertEqual(result[0]["metadata"]["parent_id"], result[1]["metadata"]["parent_id"])
        self.assertNotEqual(result[1]["metadata"]["parent_id"], result[2]["metadata"]["parent_id"])

    def test_word_tables_keep_original_block_order(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "order.docx"
            document = WordDocument()
            document.add_paragraph("表格之前")
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "字段"
            table.rows[0].cells[1].text = "值"
            document.add_paragraph("表格之后")
            document.save(path)
            result = DocumentParser().parse(path, "docx", "general")
            self.assertLess(result.text.index("表格之前"), result.text.index("【Word表格 1】"))
            self.assertLess(result.text.index("【Word表格 1】"), result.text.index("表格之后"))
    def test_query_rewrite_heuristic(self):
        rewriter = QueryRewriter()
        history = [{"role": "user", "content": "介绍朱超乐"}, {"role": "assistant", "content": "..."}]
        self.assertTrue(rewriter.should_rewrite("那他的项目呢？", history))
        self.assertFalse(rewriter.should_rewrite("请介绍公司的完整采购审批制度和适用范围", history))


if __name__ == "__main__":
    unittest.main()